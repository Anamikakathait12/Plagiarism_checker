import re
import difflib
import PyPDF2
import docx


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(file) -> str:
    """Extract text from an uploaded Flask file object (PDF, DOCX, TXT)."""
    filename = file.filename.lower()
    text = ""
    try:
        if filename.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.endswith('.txt'):
            text = file.read().decode('utf-8', errors='ignore').strip()
        else:
            return "Unsupported file format. Please upload PDF, DOCX, or TXT."
    except Exception as e:
        print(f"ERROR READING {filename}: {e}")
        return ""

    return _strip_references(text)


def extract_text_from_path(filepath: str) -> str:
    """Extract text from a saved file path on disk (PDF, DOCX, TXT)."""
    text = ""
    try:
        if filepath.endswith('.pdf'):
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        elif filepath.endswith('.docx'):
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
    except Exception as e:
        print(f"ERROR READING {filepath}: {e}")
        return ""

    return _strip_references(text)


def _strip_references(text: str) -> str:
    """Remove everything from a References/Bibliography heading onwards."""
    match = re.search(r'\n(?i)(References|Bibliography|Works Cited)\s*\n', text)
    if match:
        text = text[:match.start()]
    return text.strip()


# ---------------------------------------------------------------------------
# Side-by-side highlighting
# ---------------------------------------------------------------------------

def get_highlighted_texts(text1: str, text2: str):
    """
    Compare two texts sentence by sentence.
    Returns (html_text1, html_text2) with matching sentences wrapped in <mark>.
    """
    sentences1 = [s for s in re.split(r'(?<=[.!?]) +|\n+', text1) if len(s.strip()) > 15]
    sentences2 = [s for s in re.split(r'(?<=[.!?]) +|\n+', text2) if len(s.strip()) > 15]

    h1, h2 = text1, text2
    for s1 in sentences1:
        for s2 in sentences2:
            if difflib.SequenceMatcher(None, s1.lower(), s2.lower()).ratio() > 0.70:
                h1 = h1.replace(s1, f'<mark class="bg-danger text-white rounded px-1">{s1}</mark>')
                h2 = h2.replace(s2, f'<mark class="bg-danger text-white rounded px-1">{s2}</mark>')
    return h1, h2


# ---------------------------------------------------------------------------
# Gestalt similarity score
# ---------------------------------------------------------------------------

def get_gestalt_score(text1: str, text2: str) -> float:
    """Return a 0-100 similarity score using Gestalt pattern matching."""
    return difflib.SequenceMatcher(None, text1.lower(), text2.lower()).ratio() * 100
