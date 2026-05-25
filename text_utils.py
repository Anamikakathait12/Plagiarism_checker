import re
import difflib
import PyPDF2
import docx


def extract_text(file):
    """Extract text from an uploaded file object (PDF, DOCX, TXT)."""
    filename = file.filename.lower()
    text = ""
    try:
        if filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.endswith('.txt'):
            text = file.read().decode('utf-8', errors='ignore').strip()
        else:
            return "Unsupported file format. Please upload PDF, DOCX, or TXT."
    except Exception as e:
        print(f"\n ❌  ERROR READING {filename}: {str(e)}\n")
        return ""

    text = text.strip()
    match = re.search(r'\n(?i)(References|Bibliography|Works Cited)\s*\n', text)
    if match:
        text = text[:match.start()]
    return text.strip()


def extract_and_sanitize_text(filepath):
    """Extract and sanitize text from a file path on disk."""
    text = ""
    if filepath.endswith('.pdf'):
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                if page.extract_text():
                    text += page.extract_text() + "\n"
    elif filepath.endswith('.docx'):
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

    match = re.search(r'\n(?i)(References|Bibliography|Works Cited)\s*\n', text)
    if match:
        text = text[:match.start()]
    return text


def get_highlighted_texts(text1, text2):
    """Return HTML-highlighted versions of two texts marking similar sentences."""
    sentences1 = [s for s in re.split(r'(?<=[.!?]) +|\n+', text1) if len(s.strip()) > 15]
    sentences2 = [s for s in re.split(r'(?<=[.!?]) +|\n+', text2) if len(s.strip()) > 15]
    h1, h2 = text1, text2
    for s1 in sentences1:
        for s2 in sentences2:
            if difflib.SequenceMatcher(None, s1.lower(), s2.lower()).ratio() > 0.70:
                h1 = h1.replace(s1, f'<mark class="bg-danger text-white rounded px-1">{s1}</mark>')
                h2 = h2.replace(s2, f'<mark class="bg-danger text-white rounded px-1">{s2}</mark>')
    return h1, h2
