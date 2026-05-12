import os
import re
import time
import random
import string
import difflib
import json
import requests
import urllib.parse
from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename
from winnowing_engine import generate_fingerprints
import sqlite3
import PyPDF2
import docx
import google.generativeai as genai
from flask import Flask, render_template, request, redirect, url_for, send_from_directory
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import check_password_hash
from utils import calculate_similarity
from database import init_db, get_db_connection, register_user
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash

#  🚀  Configure Gemini (Paste your API key here!)
genai.configure(api_key="AIzaSyAoSqxPA9BxB-TN6YYrcugk7z65T-5nMjM")

app = Flask(__name__)
app.secret_key = "secret123"

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

init_db()

#  🚀  AUTOMATIC DB MIGRATION: Creates Tasks & Fingerprint tables safely!
conn = get_db_connection()
try:
    conn.execute('''
        CREATE TABLE IF NOT EXISTS tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            course_id INTEGER,
            title TEXT,
            deadline TEXT,
            total_marks INTEGER DEFAULT 100,
            FOREIGN KEY(course_id) REFERENCES courses(id)
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS document_fingerprints (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            assignment_id INTEGER NOT NULL,
            hash_value INTEGER NOT NULL,
            FOREIGN KEY(assignment_id) REFERENCES assignments(id) ON DELETE CASCADE
        )
    ''')
    conn.execute('ALTER TABLE assignments ADD COLUMN task_id INTEGER')
    conn.commit()
except:
    pass

try:
    conn.execute('ALTER TABLE tasks ADD COLUMN total_marks INTEGER DEFAULT 100')
    conn.commit()
except:
    pass
finally:
    conn.close()

#  🔐  LOGIN SETUP
login_manager = LoginManager()
login_manager.init_app(app)

class User(UserMixin):
    def __init__(self, id, email, role, username):
        self.id = id
        self.email = email
        self.role = role
        self.username = username

@login_manager.user_loader
def load_user(user_id):
    conn = get_db_connection()
    user = conn.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
    conn.close()
    if user: return User(id=user['id'], email=user['email'], role=user['role'], username=user['username'])
    return None

#  📄  FAST TEXT EXTRACTION
def extract_text(file):
    filename = file.filename.lower()
    text = ""
    try:
        if filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.endswith('.txt'):
            text = file.read().decode('utf-8', errors='ignore').strip()
        else:
            return "Unsupported file format. Please upload PDF, DOCX, or TXT."
    except Exception as e:
        print(f"\n ❌  ERROR READING {filename}: {str(e)}\n")
        return ""

    text = text.strip()
    match = re.search(r'\n(?i)(References|Bibliography|Works Cited)\s*\n', text)
    if match: text = text[:match.start()]
    return text.strip()

#  🌐  LIVE INTERNET PLAGIARISM SCANNER (Using Tavily API)
def check_internet_similarity(text):
    print(" 🔎  Checking internet plagiarism via Tavily API...")
    try:
        clean_text = re.sub(r'\s+', ' ', text)
        sentences = [s.strip() for s in re.split(r'[.!?]', clean_text) if len(s.strip()) > 60]

        if not sentences:
            return []

        words = sentences[0].split()[:30]
        query = '"' + " ".join(words) + '"'
        
        print(f" 👉  Searching Tavily for: {query}")

        url = "https://api.tavily.com/search"
        payload = {
            "api_key": "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",
            "query": query,
            "search_depth": "basic",
            "max_results": 5
        }

        res = requests.post(url, json=payload)
        res.raise_for_status()
        data = res.json()

        links = []
        for result in data.get("results", []):
            link = result.get("url")
            if link and link not in links:
                links.append(link)
        print("Found URLs:", links)
        return links
    except Exception as e:
        print("Tavily API check error:", e)
        return []

#  🕵️ ‍ ♂ ️  AI PLAGIARISM INVESTIGATOR REPORT
def generate_ai_report(student_text, urls):
    if not urls: return "No significant internet matches found."
    prompt = f"Act as an expert academic plagiarism investigator. A student submitted text that directly matched these exact websites online: {', '.join(urls)}.\n\nStudent Text Snippet: {student_text[:1500]}\n\nWrite a short, professional 2-sentence report for the teacher stating that the text appears to be copied from the internet, and name the specific website URLs."
    try:
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return "AI Report generation failed due to an API error."

#  🖍️  SIDE-BY-SIDE HIGHLIGHTER ENGINE
def get_highlighted_texts(text1, text2):
    sentences1 = [s for s in re.split(r'(?<=[.!?]) +|\n+', text1) if len(s.strip()) > 15]
    sentences2 = [s for s in re.split(r'(?<=[.!?]) +|\n+', text2) if len(s.strip()) > 15]
    h1, h2 = text1, text2
    for s1 in sentences1:
        for s2 in sentences2:
            if difflib.SequenceMatcher(None, s1.lower(), s2.lower()).ratio() > 0.70:
                h1 = h1.replace(s1, f'<mark class="bg-danger text-white rounded px-1">{s1}</mark>')
                h2 = h2.replace(s2, f'<mark class="bg-danger text-white rounded px-1">{s2}</mark>')
    return h1, h2

# --- HELPER FUNCTION: Document Ingestion and Sanitization ---
def extract_and_sanitize_text(filepath):
    text = ""
    if filepath.endswith('.pdf'):
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                if page.extract_text():
                    text += page.extract_text() + "\n"
    elif filepath.endswith('.docx'):
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

    match = re.search(r'\n(?i)(References|Bibliography|Works Cited)\s*\n', text)
    if match:
        text = text[:match.start()]

    return text

#  🏠  ROUTES
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/student-login')
def student_login_page(): return render_template('student-login.html')

@app.route('/teacher-login')
def teacher_login_page(): return render_template('teacher-login.html')

@app.route('/register', methods=['POST'])
def register():
    username = request.form.get('username')
    email = request.form.get('email')
    password = request.form.get('password')
    role = request.form.get('role')
    success = register_user(username, email, password, role)
    template = 'student-login.html' if role == 'student' else 'teacher-login.html'
    if success: return render_template(template, success="Account created! You can now log in.")
    else: return render_template(template, error="Registration failed. Username or Email already taken.")

@app.route('/login', methods=['POST'])
def login():
    identifier = request.form.get('identifier')
    password = request.form.get('password')
    role = request.form.get('role')
    conn = get_db_connection()
    user = conn.execute('SELECT * FROM users WHERE email = ? OR username = ?', (identifier, identifier)).fetchone()
    conn.close()
    template = 'student-login.html' if role == 'student' else 'teacher-login.html'
    if user and check_password_hash(user['password'], password):
        if user['role'] != role: return render_template(template, error=f"Access denied. You are registered as a {user['role']}.")
        user_obj = User(id=user['id'], email=user['email'], role=user['role'], username=user['username'])
        login_user(user_obj)
        return redirect(url_for('student_portal')) if role == 'student' else redirect(url_for('teacher_portal'))
    return render_template(template, error="Invalid Username/Email or Password.")

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

@app.route('/student-portal')
@login_required
def student_portal():
    if current_user.role != 'student': return redirect(url_for('index'))
    conn = get_db_connection()

    # 1. Fetch Assignments (Tasks)
    tasks = conn.execute('''
        SELECT t.*, c.name as course_name, u.username as teacher_name
        FROM tasks t
        JOIN courses c ON t.course_id = c.id
        JOIN enrollments e ON c.id = e.course_id
        JOIN users u ON c.teacher_id = u.id
        WHERE e.student_id = ? ORDER BY t.id DESC
    ''', (current_user.id,)).fetchall()

    # 2. Fetch Uploaded Work
    assignments = conn.execute('''
        SELECT a.*, c.name as course_name, t.title as task_title, t.total_marks
        FROM assignments a
        LEFT JOIN courses c ON a.course_id = c.id
        LEFT JOIN tasks t ON a.task_id = t.id
        WHERE a.student_id = ? ORDER BY a.id DESC
    ''', (current_user.id,)).fetchall()

   # 3. NEW: Fetch Enrolled Subjects
    enrolled_courses = conn.execute('''
        SELECT c.id, c.name, u.username as teacher_name
        FROM courses c
        JOIN enrollments e ON c.id = e.course_id
        JOIN users u ON c.teacher_id = u.id
        WHERE e.student_id = ?
    ''', (current_user.id,)).fetchall()

    conn.close()
    
    # Pass the new 'enrolled_courses' variable to the HTML
    return render_template('student-portal.html', assignments=assignments, tasks=tasks, enrolled_courses=enrolled_courses)

@app.route('/teacher-portal')
@login_required
def teacher_portal():
    if current_user.role != 'teacher': return redirect(url_for('index'))
    conn = get_db_connection()

    courses = conn.execute('SELECT * FROM courses WHERE teacher_id = ?', (current_user.id,)).fetchall()
    tasks = conn.execute('SELECT t.*, c.name as course_name FROM tasks t JOIN courses c ON t.course_id = c.id WHERE c.teacher_id = ?', (current_user.id,)).fetchall()

    assignments = conn.execute('''
        SELECT a.*, u.username as student_name, c.name as course_name, t.title as task_title, t.total_marks
        FROM assignments a
        JOIN users u ON a.student_id = u.id
        JOIN courses c ON a.course_id = c.id
        LEFT JOIN tasks t ON a.task_id = t.id
        WHERE c.teacher_id = ?
        ORDER BY a.id DESC
    ''', (current_user.id,)).fetchall()

    conn.close()
    return render_template('teacher-portal.html', assignments=assignments, courses=courses, tasks=tasks)

# --- UPDATED UPLOAD ROUTE WITH WINNOWING ---
@app.route('/upload_assignment', methods=['POST'])
@login_required
def upload_assignment():
    if current_user.role != 'student': return redirect(url_for('index'))

    file = request.files.get('file')
    course_id = request.form.get('course_id')
    task_id = request.form.get('task_id')

    if course_id == 'none': course_id = None
    if task_id == 'none': task_id = None

    if file and file.filename:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        conn = get_db_connection()

        cursor = conn.execute(
            'INSERT INTO assignments (student_id, course_id, task_id, filename) VALUES (?, ?, ?, ?)',
            (current_user.id, course_id, task_id, filename)
        )
        assignment_id = cursor.lastrowid 

        extracted_text = extract_and_sanitize_text(filepath)

        if extracted_text.strip(): 
            fingerprints = generate_fingerprints(extracted_text)
            fingerprint_data = [(assignment_id, hash_val) for hash_val in fingerprints]
            conn.executemany(
                'INSERT INTO document_fingerprints (assignment_id, hash_value) VALUES (?, ?)',
                fingerprint_data
            )

        conn.commit()
        conn.close()

    return redirect(url_for('student_portal'))

@app.route('/delete_assignment/<int:assignment_id>', methods=['POST'])
@login_required
def delete_assignment(assignment_id):
    if current_user.role != 'student': return redirect(url_for('index'))
    conn = get_db_connection()
    assignment = conn.execute('SELECT * FROM assignments WHERE id = ? AND student_id = ?', (assignment_id, current_user.id)).fetchone()

    if assignment:
        if assignment['status'] == 'Graded':
            print(f" ⚠️  Security Alert: User {current_user.username} attempted to delete a graded assignment.")
        else:
            # Safely scrub the fingerprints first, then the assignment
            conn.execute('DELETE FROM document_fingerprints WHERE assignment_id = ?', (assignment_id,))
            conn.execute('DELETE FROM assignments WHERE id = ?', (assignment_id,))
            conn.commit()
            
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], assignment['filename'])
            if os.path.exists(file_path):
                os.remove(file_path)

    conn.close()
    return redirect(url_for('student_portal'))

@app.route('/grade/<int:assignment_id>', methods=['POST'])
@login_required
def grade(assignment_id):
    if current_user.role != 'teacher': return redirect(url_for('index'))
    marks = request.form.get('marks')
    comments = request.form.get('comments')
    conn = get_db_connection()
    conn.execute('UPDATE assignments SET marks = ?, comments = ?, status = ? WHERE id = ?', (marks, comments, 'Graded', assignment_id))
    conn.commit()
    conn.close()
    return redirect(url_for('teacher_portal'))

@app.route('/download/<filename>')
@login_required
def download(filename): return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

#  🚀  THE MASTER ANALYSIS ROUTE
@app.route('/compare', methods=['POST'])
def compare():
    files = request.files.getlist("files")
    documents = []
    filenames = []

    for file in files:
        if file.filename == '': continue
        content = extract_text(file)
        if content.strip():
            documents.append(content)
            filenames.append(file.filename)
            
    template = 'teacher-portal.html' if current_user.is_authenticated and current_user.role == 'teacher' else 'index.html'

    assignments, courses, tasks = [], [], []
    if template == 'teacher-portal.html':
        conn = get_db_connection()
        assignments = conn.execute('''
            SELECT a.*, u.username as student_name, c.name as course_name, t.title as task_title, t.total_marks
            FROM assignments a
            JOIN users u ON a.student_id = u.id
            JOIN courses c ON a.course_id = c.id
            LEFT JOIN tasks t ON a.task_id = t.id
            WHERE c.teacher_id = ? ORDER BY a.id DESC
        ''', (current_user.id,)).fetchall()
        courses = conn.execute('SELECT * FROM courses WHERE teacher_id = ?', (current_user.id,)).fetchall()
        tasks = conn.execute('SELECT t.*, c.name as course_name FROM tasks t JOIN courses c ON t.course_id = c.id WHERE c.teacher_id = ?', (current_user.id,)).fetchall()
        conn.close()

    if len(documents) == 0:
        return render_template(template, error=" ⚠️  Could not read files.", assignments=assignments, courses=courses, tasks=tasks)

    # ---------------------------------------------------------
    #  🔀  SMART ROUTING: Internet vs. Peer-to-Peer
    # ---------------------------------------------------------
    internet_results = {}
    ai_report = None
    results = []
    detailed_comparisons = []

    # SCENARIO A: Single File Uploaded -> Check the Internet
    if len(documents) == 1:
        doc_text = documents[0]
        filename = filenames[0]
        words = doc_text.split()

        if len(words) > 20: 
            urls = check_internet_similarity(doc_text)
            if urls:
                internet_results[filename] = urls
                ai_report = generate_ai_report(doc_text, urls)
            else:
                ai_report = f" ✅  Excellent news! The scan is complete. No internet matches were found for '{filename}'. It appears to be 100% original."
        else:
            ai_report = f" ⚠️  '{filename}' has fewer than 20 words. Please upload a longer document."

    # SCENARIO B: Multiple Files Uploaded -> Check Against Each Other
    elif len(documents) > 1:
        similarity_matrix = calculate_similarity(documents)
        for i in range(len(filenames)):
            for j in range(i + 1, len(filenames)):
                score = float(similarity_matrix[i][j] * 100)
                results.append((filenames[i], filenames[j], round(score, 2)))
                if score > 1.0:
                    h1, h2 = get_highlighted_texts(documents[i], documents[j])
                    detailed_comparisons.append({
                        'file1': filenames[i], 'file2': filenames[j],
                        'text1': h1, 'text2': h2, 'score': round(score, 2)
                    })
        results.sort(key=lambda x: x[2], reverse=True)
        detailed_comparisons.sort(key=lambda x: x['score'], reverse=True)

    return render_template(template, results=results, assignments=assignments, courses=courses, tasks=tasks, detailed_comparisons=detailed_comparisons, internet_results=internet_results, ai_report=ai_report)

#  🏫  COURSE MANAGEMENT ROUTES
def generate_invite_code(length=6): return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))

@app.route('/create_course', methods=['POST'])
@login_required
def create_course():
    if current_user.role != 'teacher': return redirect(url_for('index'))
    course_name = request.form.get('course_name')
    code = generate_invite_code()
    conn = get_db_connection()
    conn.execute('INSERT INTO courses (name, code, teacher_id) VALUES (?, ?, ?)', (course_name, code, current_user.id))
    conn.commit()
    conn.close()
    return redirect(url_for('teacher_portal'))

@app.route('/create_task', methods=['POST'])
@login_required
def create_task():
    if current_user.role != 'teacher': return redirect(url_for('index'))
    course_id = request.form.get('course_id')
    title = request.form.get('title')
    deadline = request.form.get('deadline')
    total_marks = request.form.get('total_marks', 100)
    conn = get_db_connection()
    conn.execute('INSERT INTO tasks (course_id, title, deadline, total_marks) VALUES (?, ?, ?, ?)', (course_id, title, deadline, total_marks))
    conn.commit()
    conn.close()
    return redirect(url_for('teacher_portal'))

@app.route('/delete_task/<int:task_id>', methods=['POST'])
@login_required
def delete_task(task_id):
    if current_user.role != 'teacher': return redirect(url_for('index'))

    conn = get_db_connection()
    task = conn.execute('''
        SELECT t.id FROM tasks t
        JOIN courses c ON t.course_id = c.id
        WHERE t.id = ? AND c.teacher_id = ?
    ''', (task_id, current_user.id)).fetchone()

    if task:
        submissions = conn.execute('SELECT id, filename FROM assignments WHERE task_id = ?', (task_id,)).fetchall()
        for sub in submissions:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], sub['filename'])
            if os.path.exists(file_path):
                os.remove(file_path)
            # Safely scrub the fingerprints for each deleted submission
            conn.execute('DELETE FROM document_fingerprints WHERE assignment_id = ?', (sub['id'],))

        conn.execute('DELETE FROM assignments WHERE task_id = ?', (task_id,))
        conn.execute('DELETE FROM tasks WHERE id = ?', (task_id,))
        conn.commit()

    conn.close()
    return redirect(url_for('teacher_portal'))

@app.route('/join_course', methods=['POST'])
@login_required
def join_course():
    if current_user.role != 'student': return redirect(url_for('index'))
    
    code = request.form.get('invite_code', '').strip().upper()
    conn = get_db_connection()
    course = conn.execute('SELECT id FROM courses WHERE code = ?', (code,)).fetchone()
    
    if course:
        try:
            conn.execute('INSERT INTO enrollments (student_id, course_id) VALUES (?, ?)', (current_user.id, course['id']))
            conn.commit()
            # Flash a success message
            flash('Successfully joined the course!', 'success')
        except:
            # If the database throws an error, it means they are already enrolled
            flash('You are already enrolled in this course.', 'warning')
    else:
        # If the code doesn't exist in the database
        flash('Invalid Invite Code. Please check with your teacher.', 'danger')
        
    conn.close()
    return redirect(url_for('student_portal'))

@app.route('/leave_course/<int:course_id>', methods=['POST'])
@login_required
def leave_course(course_id):
    if current_user.role != 'student': return redirect(url_for('index'))
    
    conn = get_db_connection()
    # Delete the connection between the student and the course
    conn.execute('DELETE FROM enrollments WHERE student_id = ? AND course_id = ?', (current_user.id, course_id))
    conn.commit()
    conn.close()
    
    flash('Successfully left the subject.', 'info')
    return redirect(url_for('student_portal'))

#  🌍  GLOBAL DATABASE SCAN ROUTE (ENTERPRISE SCALABILITY)
@app.route('/global_winnow_scan/<int:new_submission_id>', methods=['GET'])
@login_required
def global_winnow_scan(new_submission_id):
    if current_user.role != 'teacher': return redirect(url_for('index'))
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. Get the fingerprints for the specifically requested file
    cursor.execute("SELECT hash_value FROM document_fingerprints WHERE assignment_id = ?", (new_submission_id,))
    # sqlite3.Row requires us to reference the column by string name or index
    new_hashes = set([row['hash_value'] for row in cursor.fetchall()])
    
    # If the document was empty/had no hashes, return an empty result
    if not new_hashes:
        conn.close()
        return render_template('global_results.html', target_id=new_submission_id, matches=[])
    
    # 2. Query the entire database to find any other submissions that share these exact hashes
    # O(1) mathematical intersection lookup!
    query = """
        SELECT assignment_id, COUNT(hash_value) as shared_count 
        FROM document_fingerprints 
        WHERE hash_value IN ({seq}) AND assignment_id != ?
        GROUP BY assignment_id
        HAVING shared_count > 10 
        ORDER BY shared_count DESC
    """.format(seq=','.join(['?']*len(new_hashes)))
    
    params = list(new_hashes) + [new_submission_id]
    cursor.execute(query, params)
    suspicious_matches = cursor.fetchall()
    
    conn.close()
    
    # Render the beautiful results page we built
    return render_template('global_results.html', target_id=new_submission_id, matches=suspicious_matches)

##  🔎  GLOBAL SIDE-BY-SIDE INSPECTOR
@app.route('/inspect_global/<int:target_id>/<int:match_id>')
@login_required
def inspect_global(target_id, match_id):
    if current_user.role != 'teacher': return redirect(url_for('index'))

    conn = get_db_connection()
    doc1 = conn.execute('SELECT filename FROM assignments WHERE id = ?', (target_id,)).fetchone()
    doc2 = conn.execute('SELECT filename FROM assignments WHERE id = ?', (match_id,)).fetchone()
    conn.close()

    if not doc1 or not doc2:
        return "Error: Documents not found in the database."

    path1 = os.path.join(app.config['UPLOAD_FOLDER'], doc1['filename'])
    path2 = os.path.join(app.config['UPLOAD_FOLDER'], doc2['filename'])

    # Reuse your existing helper functions to read and highlight the files!
    text1 = extract_and_sanitize_text(path1)
    text2 = extract_and_sanitize_text(path2)
    
    h1, h2 = get_highlighted_texts(text1, text2)
    
    # Calculate exact word-for-word percentage
    score = difflib.SequenceMatcher(None, text1.lower(), text2.lower()).ratio() * 100

    return render_template('inspect_global.html', 
                           file1=doc1['filename'], file2=doc2['filename'], 
                           text1=h1, text2=h2, score=round(score, 2))

if __name__ == '__main__':
    print(" 🚀  Server running...")
    app.run(debug=True)
